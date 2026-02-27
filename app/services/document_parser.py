import asyncio
import logging
import re
from io import BytesIO
from pathlib import Path
from typing import Tuple

import pandas as pd
import pdfplumber
from docx import Document
from fastapi import UploadFile

logger = logging.getLogger(__name__)

MAX_FILE_SIZE_MB = 10
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
MAX_OUTPUT_CHARS = 12000
PARSE_TIMEOUT_SECONDS = 30


async def extract_text_from_upload(upload_file: UploadFile) -> str:
    """
    Extract clean text from an uploaded business document.

    Supported formats:
    - .pdf
    - .docx
    - .xlsx / .xls
    """
    filename = upload_file.filename or ""
    extension = Path(filename).suffix.lower()

    content = await upload_file.read()
    if not content:
        raise ValueError("Uploaded file is empty")

    if len(content) > MAX_FILE_SIZE_BYTES:
        raise ValueError(f"File exceeds {MAX_FILE_SIZE_MB}MB limit")

    if extension == ".docx":
        text = await asyncio.wait_for(
            asyncio.to_thread(_extract_docx_text, content),
            timeout=PARSE_TIMEOUT_SECONDS,
        )
    elif extension == ".pdf":
        text = await asyncio.wait_for(
            asyncio.to_thread(_extract_pdf_text, content),
            timeout=PARSE_TIMEOUT_SECONDS,
        )
    elif extension in {".xlsx", ".xls"}:
        text = await asyncio.wait_for(
            asyncio.to_thread(_extract_excel_text, content, extension),
            timeout=PARSE_TIMEOUT_SECONDS,
        )
    else:
        raise ValueError("Unsupported file type. Allowed: .pdf, .docx, .xlsx, .xls")

    raw_len = len(text)
    cleaned = _clean_text(text)
    if not cleaned:
        raise ValueError("No readable text extracted from uploaded file")

    truncated = False
    if len(cleaned) > MAX_OUTPUT_CHARS:
        cleaned = cleaned[:MAX_OUTPUT_CHARS] + "\n\n[... content truncated ...]"
        truncated = True

    logger.info(
        "Document parsed: filename=%s extension=%s raw_chars=%d cleaned_chars=%d truncated=%s",
        filename,
        extension,
        raw_len,
        len(cleaned),
        truncated,
    )
    logger.debug("Document content preview (first 1000 chars): %s", cleaned[:1000])
    
    return cleaned


def _extract_docx_text(content: bytes) -> str:
    """
    Extract text from DOCX including paragraphs, tables, and headers/footers.
    
    Tables are formatted as markdown-style pipe-separated rows.
    """
    try:
        doc = Document(BytesIO(content))
        chunks: list[str] = []
        
        for section in doc.sections:
            header_text = _extract_docx_header_footer(section.header)
            if header_text:
                chunks.append(f"[Header]\n{header_text}")
        
        for element in doc.element.body:
            if element.tag.endswith('}p'):
                for para in doc.paragraphs:
                    if para._element is element:
                        text = para.text.strip()
                        if text:
                            chunks.append(text)
                        break
            elif element.tag.endswith('}tbl'):
                for table in doc.tables:
                    if table._element is element:
                        table_text = _extract_docx_table(table)
                        if table_text:
                            chunks.append(table_text)
                        break
        
        for section in doc.sections:
            footer_text = _extract_docx_header_footer(section.footer)
            if footer_text:
                chunks.append(f"[Footer]\n{footer_text}")
        
        result = "\n\n".join(chunks)
        logger.debug("DOCX extraction complete: elements=%d chars=%d", len(chunks), len(result))
        return result
        
    except Exception as exc:
        logger.exception("DOCX parsing failed")
        raise ValueError(f"Failed to parse DOCX: {exc}") from exc


def _extract_docx_table(table) -> str:
    """Extract text from a DOCX table as pipe-separated rows."""
    lines: list[str] = []
    
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        if any(cells):
            lines.append(" | ".join(cells))
    
    return "\n".join(lines) if lines else ""


def _extract_docx_header_footer(hf) -> str:
    """Extract text from a DOCX header or footer."""
    if hf is None:
        return ""
    
    paragraphs = [p.text.strip() for p in hf.paragraphs if p.text and p.text.strip()]
    return "\n".join(paragraphs)


def _extract_pdf_text(content: bytes) -> str:
    """
    Extract text from PDF using pdfplumber with table support.
    
    Tables are formatted as markdown-style pipe-separated rows.
    Regular text is extracted with better layout awareness.
    """
    try:
        chunks: list[str] = []
        
        with pdfplumber.open(BytesIO(content)) as pdf:
            for idx, page in enumerate(pdf.pages):
                try:
                    page_chunks: list[str] = []
                    
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            table_text = _format_table_as_text(table)
                            if table_text:
                                page_chunks.append(table_text)
                    
                    page_text = page.extract_text(layout=True) or ""
                    page_text = page_text.strip()
                    if page_text:
                        page_chunks.append(page_text)
                    
                    if page_chunks:
                        chunks.append("\n\n".join(page_chunks))
                        
                except Exception as e:
                    logger.warning("Skipping unreadable PDF page index=%d error=%s", idx, str(e))
        
        result = "\n\n".join(chunks)
        logger.debug("PDF extraction complete: pages=%d chars=%d", len(chunks), len(result))
        return result
        
    except Exception as exc:
        logger.exception("PDF parsing failed")
        raise ValueError(f"Failed to parse PDF: {exc}") from exc


def _format_table_as_text(table: list) -> str:
    """
    Format a table (list of rows) as markdown-style pipe-separated text.
    """
    if not table:
        return ""
    
    lines: list[str] = []
    for row in table:
        if row is None:
            continue
        cells = [str(cell).strip() if cell else "" for cell in row]
        if any(cells):
            lines.append(" | ".join(cells))
    
    return "\n".join(lines) if lines else ""


def _extract_excel_text(content: bytes, extension: str) -> str:
    try:
        engine = "openpyxl" if extension == ".xlsx" else "xlrd"
        sheets = pd.read_excel(BytesIO(content), sheet_name=None, engine=engine)

        lines: list[str] = []
        for sheet_name, df in sheets.items():
            if df is None or df.empty:
                continue
            lines.append(f"Sheet: {sheet_name}")
            for row in df.fillna("").astype(str).values.tolist():
                row_text = " | ".join(cell.strip() for cell in row if cell.strip())
                if row_text:
                    lines.append(row_text)
            lines.append("")
        return "\n".join(lines)
    except Exception as exc:
        logger.exception("Excel parsing failed")
        raise ValueError(f"Failed to parse Excel: {exc}") from exc


def _clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()
